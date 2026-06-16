"""
app/routers/tool_export_word.py

POST /api/tools/export-word
接收 Markdown 报告文本，生成 Word 文档并返回文件流。
"""
import io
import re
from datetime import datetime
from urllib.parse import quote

import pytz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

from app.middlewares.auth import require_password_changed
from app.models.user import User

router = APIRouter(prefix="/tools", tags=["tools"])

_FONT = "微软雅黑"
_BODY_SIZE = 22  # half-points = 11pt


def _set_font(run, size_hpt: int | None = None) -> None:
    run.font.name = _FONT
    run.element.rPr.rFonts.set(qn("w:eastAsia"), _FONT)
    if size_hpt is not None:
        run.font.size = Pt(size_hpt / 2)


def _parse_inline(para, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        _set_font(run, _BODY_SIZE)


def _markdown_to_doc(doc: Document, content: str) -> None:
    for line in content.split("\n"):
        stripped = line.rstrip()

        if stripped == "":
            doc.add_paragraph()
            continue

        h_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            spacings = {1: (240, 120), 2: (200, 100), 3: (160, 80)}
            before, after = spacings[level]
            p = doc.add_heading(h_match.group(2), level=level)
            p.paragraph_format.space_before = Pt(before / 20)
            p.paragraph_format.space_after = Pt(after / 20)
            continue

        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            try:
                para = doc.add_paragraph(style="List Bullet")
            except KeyError:
                para = doc.add_paragraph()
            _parse_inline(para, text)
            continue

        if stripped.startswith("> "):
            text = stripped[2:]
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            run = para.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            _set_font(run, _BODY_SIZE)
            continue

        para = doc.add_paragraph()
        _parse_inline(para, stripped)


class ExportWordRequest(BaseModel):
    content: str
    title: str = "千川剪辑预审报告"


@router.post("/export-word")
async def export_word(
    body: ExportWordRequest,
    current_user: User = Depends(require_password_changed),
):
    """Word 导出接口：Markdown → .docx 文件流。"""
    if not body.content.strip():
        raise HTTPException(
            status_code=400,
            detail={"code": "INVALID_INPUT", "message": "内容为空"},
        )

    doc = Document()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(body.title)
    title_run.bold = True
    _set_font(title_run, 32)

    tz = pytz.timezone("Asia/Shanghai")
    now_str = datetime.now(tz).strftime("%Y-%m-%d %H:%M:%S")
    time_para = doc.add_paragraph()
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_run = time_para.add_run(f"导出时间：{now_str}")
    _set_font(time_run, 20)

    doc.add_paragraph()

    _markdown_to_doc(doc, body.content)

    buf = io.BytesIO()
    doc.save(buf)
    file_bytes = buf.getvalue()

    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"千川预审报告_{date_str}.docx"
    encoded = quote(filename)

    return Response(
        content=file_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded}"},
    )
