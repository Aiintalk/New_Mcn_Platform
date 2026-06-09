"""
app/services/intake_report.py

红人入驻问卷报告文件生成：
- generate_docx(submission_id, ai_report, kol_name) → storage/intake_reports/{id}.docx
- generate_pdf(submission_id, ai_report, kol_name)  → storage/intake_reports/{id}.pdf

依赖：python-docx（docx）、reportlab（pdf）
"""
import os
import re
from datetime import datetime, timezone
from pathlib import Path

REPORT_DIR = Path("storage/intake_reports")


def _ensure_dir() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)


def _strip_markdown(text: str) -> str:
    """将 Markdown 简单转换为纯文本（用于 Word 段落）。"""
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", text)
    return text.strip()


def generate_docx(submission_id: int | str, ai_report: str, kol_name: str | None = None) -> str:
    """
    生成 Word 报告，返回文件路径（相对路径）。
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ensure_dir()

    doc = Document()

    # 标题
    title = doc.add_heading(f"新红人分析报告 · {kol_name or '未知'}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 分割线 & 元数据
    doc.add_paragraph("─" * 40)
    meta = doc.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC")

    doc.add_paragraph()

    # 逐段写入 AI 报告
    for line in ai_report.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        # Markdown 标题 → Word 标题
        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            level = min(len(h_match.group(1)), 4)
            doc.add_heading(h_match.group(2), level=level)
            continue

        # 普通段落（保留加粗）
        para = doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    # 底部签名
    doc.add_paragraph()
    doc.add_paragraph("─" * 40)
    footer = doc.add_paragraph("本报告由达人说平台自动生成")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rel_path = str(REPORT_DIR / f"{submission_id}.docx")
    doc.save(rel_path)
    return rel_path


def generate_pdf(submission_id: int | str, ai_report: str, kol_name: str | None = None) -> str:
    """
    生成 PDF 报告，返回文件路径（相对路径）。
    使用 reportlab 直接生成，兼容 Windows。
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    _ensure_dir()

    # 尝试注册中文字体（Windows 系统字体）
    font_name = "Helvetica"
    for font_path, name in [
        ("C:/Windows/Fonts/msyh.ttc",    "MsYaHei"),
        ("C:/Windows/Fonts/simsun.ttc",   "SimSun"),
        ("C:/Windows/Fonts/simhei.ttf",   "SimHei"),
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(name, font_path))
                font_name = name
                break
            except Exception:
                continue

    rel_path = str(REPORT_DIR / f"{submission_id}.pdf")
    doc = SimpleDocTemplate(
        rel_path,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "CnBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        spaceAfter=6,
    )
    h1_style = ParagraphStyle(
        "CnH1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=22,
        spaceAfter=10,
    )
    h2_style = ParagraphStyle(
        "CnH2",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=13,
        leading=20,
        spaceAfter=8,
    )

    story = []

    # 标题
    story.append(Paragraph(f"新红人分析报告 · {kol_name or '未知'}", h1_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC",
        body_style,
    ))
    story.append(Spacer(1, 0.5 * cm))

    # 逐段处理 AI 报告
    for line in ai_report.split("\n"):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.2 * cm))
            continue

        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            style = h1_style if level == 1 else h2_style
            story.append(Paragraph(h_match.group(2), style))
            continue

        # 简单加粗处理
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", stripped)
        story.append(Paragraph(text, body_style))

    story.append(Spacer(1, 0.5 * cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
    story.append(Paragraph("本报告由达人说平台自动生成", body_style))

    doc.build(story)
    return rel_path
