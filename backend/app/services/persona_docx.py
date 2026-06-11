"""
app/services/persona_docx.py

人格定位 Markdown → Word 文档生成。
"""
import os
import re
from datetime import datetime, timezone

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

STORAGE_DIR = "storage/persona_reports"

# 确保存储目录存在
os.makedirs(STORAGE_DIR, exist_ok=True)


def generate_persona_docx(
    report_id: int,
    doc_type: str,
    content: str,
    influencer_name: str,
) -> str:
    """
    将 Markdown 内容转换为 Word 文档。

    Args:
        report_id: 报告 ID
        doc_type: "profile" 或 "plan"
        content: Markdown 格式内容
        influencer_name: 达人名字

    Returns:
        文件路径
    """
    doc = Document()

    # 文档头部
    title = f"{influencer_name} · {'人格档案' if doc_type == 'profile' else '内容规划'}"
    heading = doc.add_heading(title, level=1)

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
    meta = doc.add_paragraph(f"生成时间：{date_str}")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 解析 Markdown 并写入段落
    _markdown_to_paragraphs(doc, content)

    # 保存
    filename = f"{report_id}_{doc_type}.docx"
    filepath = os.path.join(STORAGE_DIR, filename)
    doc.save(filepath)
    return filepath


def generate_questionnaire_template() -> str:
    """
    生成问卷模板 Word 文档。
    返回文件路径。
    """
    doc = Document()

    # 页面标题
    title = doc.add_heading("达人入职信息采集表", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    sections = [
        {
            "name": "一、基本信息",
            "color": RGBColor(0x7C, 0x3A, 0xED),  # 紫色主题
            "questions": [
                ("1. 达人的名字 / 昵称是什么？（粉丝怎么称呼 ta）", True),
                ("2. 年龄、所在城市？", True),
                ("3. 职业背景和从业经历？（做过什么、多少年、怎么走到今天的）", True),
                ("4. 想做的内容赛道是什么？（如美妆、母婴、美食等）", True),
            ],
        },
        {
            "name": "二、个人特色",
            "color": RGBColor(0x7C, 0x3A, 0xED),
            "questions": [
                ("5. 有什么专业资质、成就或独特经历？", True),
                ("6. 性格特点是什么？朋友会怎么形容 ta？", True),
                ("7. 想要什么样的说话风格？", True),
            ],
        },
        {
            "name": "三、内容方向",
            "color": RGBColor(0x7C, 0x3A, 0xED),
            "questions": [
                ("8. 目标受众是什么人群？", False),
                ("9. 有没有想对标或喜欢的博主？喜欢 ta 什么？", False),
                ("10. 还有什么想补充的？", False),
            ],
        },
    ]

    for section in sections:
        heading = doc.add_heading(section["name"], level=2)
        for run in heading.runs:
            run.font.color.rgb = section["color"]

        for q_text, required in section["questions"]:
            p = doc.add_paragraph()
            run = p.add_run(q_text)
            if required:
                run2 = p.add_run(" *必填")
                run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                run2.font.size = Pt(9)
            # 答案留白区域
            doc.add_paragraph("")

    filepath = os.path.join(STORAGE_DIR, "questionnaire_template.docx")
    doc.save(filepath)
    return filepath


def _markdown_to_paragraphs(doc: Document, content: str) -> None:
    """
    将 Markdown 内容转换为 Word 段落。

    支持：
    - # ## ### #### → 标题 1-4
    - - / * → 无序列表
    - 1. → 有序列表
    - > → 引用（斜体+缩进）
    - **text** → 粗体
    - 空行 → 空段落
    """
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # 引用
        if line.startswith("> "):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue

        # 无序列表
        list_match = re.match(r"^[\-\*]\s+(.+)$", line)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, text)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if ol_match:
            text = ol_match.group(2)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_inline_runs(p, line.strip())
        i += 1


def _add_inline_runs(paragraph, text: str) -> None:
    """
    解析行内格式（**粗体**），添加 run 到段落。
    """
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)
