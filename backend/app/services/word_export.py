"""
app/services/word_export.py

共用 Markdown → Word 文档生成，返回 bytes。
tiktok-writer / 其他工具复用此模块。

支持语法：
  # / ## / ### → Heading 1/2/3
  - text / * text → Bullet List
  **bold** → Bold run
  空行 → 空段落
  1. xxx → 普通段落（保留原版有序列表 bug，不修复）
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _set_run_font(run, font_name: str, font_size_pt: int | None = None) -> None:
    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def _parse_inline(para, text: str, font_name: str, font_size_pt: int | None) -> None:
    """将 **bold** 标记应用到 para 的 runs 上。"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        _set_run_font(run, font_name, font_size_pt)


def markdown_to_docx_bytes(
    title: str,
    metadata_lines: list[str],
    content: str,
    font_name: str = "Arial",
    body_font_size_pt: int = 22,
) -> bytes:
    """
    生成 Word 文档并以 bytes 返回。

    Args:
        title: 文档标题行（居中，加粗）
        metadata_lines: 元数据行列表，如 ["Topic: ...", "Exported: ..."]（居中，小字）
        content: Markdown 格式的正文内容
        font_name: 正文字体，默认 Arial
        body_font_size_pt: 正文字号（pt），默认 22
    """
    doc = Document()

    # 标题（居中加粗）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = font_name
    title_run.font.size = Pt(28)

    # 元数据行（居中，12pt）
    for meta in metadata_lines:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(meta)
        _set_run_font(mr, font_name, 12)

    # 间隔行
    doc.add_paragraph()

    # 正文：逐行解析 Markdown
    for line in content.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Heading（# / ## / ###）
        h_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2), level=level)
            continue

        # Bullet list（- / *）
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            try:
                para = doc.add_paragraph(style="List Bullet")
            except KeyError:
                para = doc.add_paragraph()
                para.add_run("• ")
            _parse_inline(para, text, font_name, body_font_size_pt)
            continue

        # 普通段落（含有序列表——保留原版 bug，渲染为普通段落）
        para = doc.add_paragraph()
        _parse_inline(para, stripped, font_name, body_font_size_pt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
