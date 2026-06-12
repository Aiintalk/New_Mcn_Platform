"""
app/services/benchmark_report.py

对标分析报告文件生成：
- generate_docx(analysis_id, content, account_name, doc_type) → storage/benchmark_reports/{id}_{type}.docx

复用 intake_report.py 的 Markdown → docx 转换逻辑。
"""
import re
from datetime import datetime, timezone
from pathlib import Path

REPORT_DIR = Path("storage/benchmark_reports")


def _ensure_dir() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)


def generate_docx(analysis_id: int | str, content: str, account_name: str | None = None, doc_type: str = "profile") -> str:
    """
    生成 Word 文档，返回文件路径（相对路径）。
    doc_type: 'profile' 或 'plan'
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ensure_dir()

    doc = Document()

    label = "人格档案" if doc_type == "profile" else "内容规划"
    title = doc.add_heading(f"{label} · {account_name or '未知'}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 40)
    meta = doc.add_paragraph()
    meta.add_run(f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} UTC")
    doc.add_paragraph()

    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue

        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            level = min(len(h_match.group(1)), 4)
            doc.add_heading(h_match.group(2), level=level)
            continue

        para = doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    doc.add_paragraph()
    doc.add_paragraph("─" * 40)
    footer = doc.add_paragraph("本报告由 MCN 平台自动生成")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rel_path = str(REPORT_DIR / f"{analysis_id}_{doc_type}.docx")
    doc.save(rel_path)
    return rel_path
