"""
Unit tests for parse_qianchuan_review_file。

与 selling-point 版本的关键区别：
- PDF 不支持（返回提示文字，不抛错）
- .pages 有日历噪声过滤（星期、月份季度、公元）
"""
import io
import zipfile
from unittest.mock import AsyncMock, MagicMock

import pytest

from app.services.file_parser import parse_qianchuan_review_file


def _mock_file(filename: str, content: bytes) -> MagicMock:
    f = MagicMock()
    f.filename = filename
    f.read = AsyncMock(return_value=content)
    return f


# ---------- txt / md ----------

@pytest.mark.asyncio
async def test_txt_returns_text():
    result = await parse_qianchuan_review_file(_mock_file("script.txt", "脚本内容第一行标题".encode()))
    assert result == "脚本内容第一行标题"


@pytest.mark.asyncio
async def test_md_returns_text():
    result = await parse_qianchuan_review_file(_mock_file("script.md", "# 标题\n内容".encode()))
    assert "标题" in result


# ---------- docx ----------

@pytest.mark.asyncio
async def test_docx_extracts_paragraphs():
    from docx import Document
    doc = Document()
    doc.add_paragraph("千川脚本第一段")
    doc.add_paragraph("第二段内容")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = await parse_qianchuan_review_file(_mock_file("script.docx", buf.read()))
    assert "千川脚本第一段" in result
    assert "第二段内容" in result


# ---------- pdf —— 不支持，返回提示 ----------

@pytest.mark.asyncio
async def test_pdf_returns_unsupported_hint():
    result = await parse_qianchuan_review_file(_mock_file("data.pdf", b"%PDF-1.4"))
    assert result == "[暂不支持 PDF 格式，请转为 .docx 或 .txt 后上传]"


# ---------- .pages —— 基础提取 ----------

def _make_pages_zip(text: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("Index/Document.iwa", b"\x00\x00\x00\x00" + text.encode())
    buf.seek(0)
    return buf.read()


@pytest.mark.asyncio
async def test_pages_extracts_chinese():
    pages_bytes = _make_pages_zip("这是一段产品脚本内容，超过十个中文字的段落。")
    result = await parse_qianchuan_review_file(_mock_file("script.pages", pages_bytes))
    assert "产品脚本内容" in result


@pytest.mark.asyncio
async def test_pages_filters_short_chinese():
    """少于5个汉字的片段应被过滤"""
    pages_bytes = _make_pages_zip("两字" + "A" * 20)
    result = await parse_qianchuan_review_file(_mock_file("noise.pages", pages_bytes))
    assert "两字" not in result


# ---------- .pages —— 日历噪声过滤（与 selling-point 的关键差异） ----------

@pytest.mark.asyncio
async def test_pages_filters_weekday_noise():
    """星期X[BJR] 模式应被过滤"""
    pages_bytes = _make_pages_zip("星期一B这是噪声")
    result = await parse_qianchuan_review_file(_mock_file("cal.pages", pages_bytes))
    assert "星期一B" not in result


@pytest.mark.asyncio
async def test_pages_filters_month_noise():
    """[一二...十]+月 开头且长度<20 的片段应被过滤"""
    pages_bytes = _make_pages_zip("一月二日三日")
    result = await parse_qianchuan_review_file(_mock_file("cal.pages", pages_bytes))
    assert "一月" not in result


@pytest.mark.asyncio
async def test_pages_filters_quarter_noise():
    """第[一二三四]季度 且长度<20 的片段应被过滤"""
    pages_bytes = _make_pages_zip("第一季度数据")
    result = await parse_qianchuan_review_file(_mock_file("cal.pages", pages_bytes))
    assert "第一季度" not in result


@pytest.mark.asyncio
async def test_pages_filters_gongyan_noise():
    """公元开头且长度<10 的片段应被过滤"""
    pages_bytes = _make_pages_zip("公元前")
    result = await parse_qianchuan_review_file(_mock_file("cal.pages", pages_bytes))
    assert "公元前" not in result


@pytest.mark.asyncio
async def test_pages_keeps_long_month_content():
    """长度>=20 的月份内容不应被过滤"""
    long_month = "一月份这个产品的卖点非常明显价格优惠超值"  # >= 20 字
    pages_bytes = _make_pages_zip(long_month)
    result = await parse_qianchuan_review_file(_mock_file("script.pages", pages_bytes))
    assert "一月份" in result


@pytest.mark.asyncio
async def test_pages_missing_iwa():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("dummy.txt", "nothing")
    buf.seek(0)
    result = await parse_qianchuan_review_file(_mock_file("empty.pages", buf.read()))
    assert "格式异常" in result


@pytest.mark.asyncio
async def test_pages_invalid_zip():
    result = await parse_qianchuan_review_file(_mock_file("bad.pages", b"not a zip"))
    assert "格式异常" in result


# ---------- 未知格式 ----------

@pytest.mark.asyncio
async def test_unknown_ext_raises_value_error():
    with pytest.raises(ValueError, match="不支持的文件格式"):
        await parse_qianchuan_review_file(_mock_file("data.xlsx", b"content"))
