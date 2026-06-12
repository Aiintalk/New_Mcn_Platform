"""
Unit tests for app/services/word_export.py
测试 Markdown → docx 转换逻辑，全部用内存，无 DB 依赖。
"""
import io
import pytest
from docx import Document


def load_doc(docx_bytes: bytes) -> Document:
    return Document(io.BytesIO(docx_bytes))


def collect_text(doc: Document) -> list[str]:
    """返回文档所有段落的文本列表（空段落记为 ''）。"""
    return [p.text for p in doc.paragraphs]


class TestMarkdownToDocxBytes:

    def test_returns_bytes(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="Test Title",
            metadata_lines=["Topic: foo", "Exported: 2026-01-01"],
            content="Hello world",
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx_format(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="My Title",
            metadata_lines=["Topic: t"],
            content="plain text",
        )
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any("My Title" in t for t in texts)

    def test_title_in_document(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="PersonaName · TikTok Script",
            metadata_lines=[],
            content="body",
        )
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any("PersonaName · TikTok Script" in t for t in texts)

    def test_metadata_lines_in_document(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=["Topic: https://example.com", "Exported: 2026-06-12"],
            content="body",
        )
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any("Topic: https://example.com" in t for t in texts)
        assert any("Exported: 2026-06-12" in t for t in texts)

    def test_heading1_becomes_heading_paragraph(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="# My Heading",
        )
        doc = load_doc(result)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("My Heading" in h.text for h in headings)

    def test_heading2_becomes_heading_paragraph(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="## Section Two",
        )
        doc = load_doc(result)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Section Two" in h.text for h in headings)

    def test_bold_text_has_bold_run(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="normal **bold word** after",
        )
        doc = load_doc(result)
        bold_runs = [
            run for p in doc.paragraphs for run in p.runs if run.bold and run.text.strip()
        ]
        assert any("bold word" in r.text for r in bold_runs)

    def test_bullet_list_item(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="- List item one",
        )
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any("List item one" in t for t in texts)

    def test_ordered_list_rendered_as_plain(self):
        """1. xxx 应渲染为普通段落（保留原版 bug，不修复）。"""
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="1. First item",
        )
        doc = load_doc(result)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert not any("First item" in h.text for h in headings)
        texts = collect_text(doc)
        assert any("First item" in t for t in texts)

    def test_empty_line_produces_empty_paragraph(self):
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="T",
            metadata_lines=[],
            content="line one\n\nline two",
        )
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any(t == "" for t in texts)
        assert any("line one" in t for t in texts)
        assert any("line two" in t for t in texts)

    def test_empty_content_produces_valid_docx(self):
        """空 content 不应报错，应返回合法 docx（只有标题行）。"""
        from app.services.word_export import markdown_to_docx_bytes
        result = markdown_to_docx_bytes(
            title="Title Only",
            metadata_lines=[],
            content="",
        )
        assert isinstance(result, bytes)
        assert len(result) > 0
        doc = load_doc(result)
        texts = collect_text(doc)
        assert any("Title Only" in t for t in texts)
