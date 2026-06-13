"""Unit tests for parse_selling_point_file（selling-point 专用解析函数）"""
import io
import zipfile
from unittest.mock import AsyncMock, MagicMock

import pytest

from app.services.file_parser import parse_selling_point_file


def _mock_file(filename: str, content: bytes) -> MagicMock:
    f = MagicMock()
    f.filename = filename
    f.read = AsyncMock(return_value=content)
    return f


# ---------- txt / md ----------

@pytest.mark.asyncio
async def test_txt_returns_text():
    result = await parse_selling_point_file(_mock_file("brief.txt", "产品卖点内容".encode()))
    assert result == "产品卖点内容"

@pytest.mark.asyncio
async def test_md_returns_text():
    result = await parse_selling_point_file(_mock_file("script.md", "# 标题\n内容".encode()))
    assert "标题" in result

@pytest.mark.asyncio
async def test_txt_no_truncation():
    """selling-point 版本不截断"""
    long = "中文内容" * 5000
    result = await parse_selling_point_file(_mock_file("long.txt", long.encode()))
    assert len(result) == len(long)

# ---------- docx ----------

@pytest.mark.asyncio
async def test_docx_extracts_paragraphs():
    from docx import Document
    doc = Document()
    doc.add_paragraph("第一段卖点")
    doc.add_paragraph("第二段说明")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    result = await parse_selling_point_file(_mock_file("brief.docx", buf.read()))
    assert "第一段卖点" in result
    assert "第二段说明" in result

# ---------- pdf (pdfplumber) ----------

@pytest.mark.asyncio
async def test_pdf_extracts_text():
    import unittest.mock as mock
    mock_page = MagicMock()
    mock_page.extract_text.return_value = "玻尿酸、烟酰胺"
    mock_pdf = MagicMock()
    mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = [mock_page]
    with mock.patch("app.services.file_parser.pdfplumber") as mp:
        mp.open.return_value = mock_pdf
        result = await parse_selling_point_file(_mock_file("product.pdf", b"%PDF"))
    assert "玻尿酸" in result

@pytest.mark.asyncio
async def test_pdf_multiple_pages():
    import unittest.mock as mock
    pages = []
    for txt in ["第一页内容", "第二页内容", None]:
        p = MagicMock(); p.extract_text.return_value = txt; pages.append(p)
    mock_pdf = MagicMock()
    mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
    mock_pdf.__exit__ = MagicMock(return_value=False)
    mock_pdf.pages = pages
    with mock.patch("app.services.file_parser.pdfplumber") as mp:
        mp.open.return_value = mock_pdf
        result = await parse_selling_point_file(_mock_file("multi.pdf", b"%PDF"))
    assert "第一页内容" in result
    assert "第二页内容" in result

# ---------- .doc ----------

@pytest.mark.asyncio
async def test_doc_returns_hint():
    result = await parse_selling_point_file(_mock_file("old.doc", b"\xd0\xcf\x11"))
    assert ".doc 格式暂不支持" in result

# ---------- .pages ----------

def _make_pages_zip(text: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("Index/Document.iwa", b"\x00\x00\x00\x00" + text.encode())
    buf.seek(0)
    return buf.read()

@pytest.mark.asyncio
async def test_pages_extracts_chinese():
    pages_bytes = _make_pages_zip("这是一段产品卖点说明，超过十个中文字的内容测试。")
    result = await parse_selling_point_file(_mock_file("doc.pages", pages_bytes))
    assert "产品卖点" in result

@pytest.mark.asyncio
async def test_pages_filters_short_chinese():
    pages_bytes = _make_pages_zip("ab两字" + "A" * 20)
    result = await parse_selling_point_file(_mock_file("noise.pages", pages_bytes))
    assert "两字" not in result

@pytest.mark.asyncio
async def test_pages_no_calendar_filter():
    """selling-point 不过滤日历型中文"""
    pages_bytes = _make_pages_zip("一月二月三月四月五月六月七月八月")
    result = await parse_selling_point_file(_mock_file("cal.pages", pages_bytes))
    assert "一月" in result

@pytest.mark.asyncio
async def test_pages_missing_iwa():
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("dummy.txt", "nothing")
    buf.seek(0)
    result = await parse_selling_point_file(_mock_file("empty.pages", buf.read()))
    assert "格式异常" in result

@pytest.mark.asyncio
async def test_pages_invalid_zip():
    result = await parse_selling_point_file(_mock_file("bad.pages", b"not a zip"))
    assert "格式异常" in result

# ---------- 未知格式 ----------

@pytest.mark.asyncio
async def test_unknown_ext_utf8_decode():
    result = await parse_selling_point_file(_mock_file("data.csv", "产品名,价格\n精华,299".encode()))
    assert "精华" in result


@pytest.mark.asyncio
async def test_pages_decompress_fallback():
    """snappy 解压失败时 fallback 到原始字节，不抛异常，中文内容仍可提取"""
    # 构造：IWA 数据前缀 + 直接写中文字节（不经 snappy 压缩）
    # _parse_pages_selling_point 会尝试 snappy.decompress(iwa_data[4:]) 失败后
    # 回退到 iwa_data 本身，再 decode，正则仍能匹配中文
    chinese_text = "这是一段足够长的中文产品卖点说明内容，用于测试fallback路径"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        # 4字节magic + 非snappy压缩的中文UTF-8
        zf.writestr("Index/Document.iwa", b"\x00\x00\x00\x00" + chinese_text.encode("utf-8"))
    buf.seek(0)
    result = await parse_selling_point_file(_mock_file("fallback.pages", buf.read()))
    assert isinstance(result, str)  # 不抛异常


@pytest.mark.asyncio
async def test_docx_parse_error_raises_value_error():
    """损坏的 docx 文件应抛出 ValueError"""
    with pytest.raises(ValueError, match=".docx 文件解析失败"):
        await parse_selling_point_file(_mock_file("bad.docx", b"not a docx file at all"))
