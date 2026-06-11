"""Unit tests for app.services.intake_report — report file generation."""
import os
from pathlib import Path

import pytest

from app.services.intake_report import _strip_markdown, generate_docx


class TestStripMarkdown:
    def test_strip_markdown_removes_h1(self):
        assert _strip_markdown("# Title") == "Title"

    def test_strip_markdown_removes_h2(self):
        assert _strip_markdown("## Section") == "Section"

    def test_strip_markdown_removes_h3(self):
        assert _strip_markdown("### Sub") == "Sub"

    def test_strip_markdown_removes_bold(self):
        assert _strip_markdown("**bold**") == "bold"

    def test_strip_markdown_removes_italic(self):
        assert _strip_markdown("*italic*") == "italic"

    def test_strip_markdown_removes_underscore_italic(self):
        assert _strip_markdown("_italic_") == "italic"

    def test_strip_markdown_mixed(self):
        result = _strip_markdown("## **Bold Title** and *italic*")
        assert "**" not in result
        assert "*" not in result or result.count("*") == 0


class TestGenerateDocx:
    def test_generate_docx_creates_file(self, tmp_path):
        output_dir = tmp_path / "reports"
        output_dir.mkdir()

        import app.services.intake_report as mod
        original_dir = mod.REPORT_DIR
        mod.REPORT_DIR = output_dir

        try:
            path = mod.generate_docx(
                submission_id="test_001",
                ai_report="# Report\n\nSome content\n\n**Bold text**",
                kol_name="测试红人",
            )
            assert Path(path).exists()
            assert path.endswith(".docx")
        finally:
            mod.REPORT_DIR = original_dir

    def test_generate_docx_contains_kol_name(self, tmp_path):
        output_dir = tmp_path / "reports"
        output_dir.mkdir()

        import app.services.intake_report as mod
        original_dir = mod.REPORT_DIR
        mod.REPORT_DIR = output_dir

        try:
            path = mod.generate_docx(
                submission_id="test_002",
                ai_report="内容",
                kol_name="小红",
            )
            from docx import Document
            doc = Document(path)
            # Check title contains kol_name
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            assert any("小红" in h for h in headings)
        finally:
            mod.REPORT_DIR = original_dir

    def test_generate_docx_handles_markdown_headings(self, tmp_path):
        output_dir = tmp_path / "reports"
        output_dir.mkdir()

        import app.services.intake_report as mod
        original_dir = mod.REPORT_DIR
        mod.REPORT_DIR = output_dir

        try:
            path = mod.generate_docx(
                submission_id="test_003",
                ai_report="# Main\n## Sub\n### Detail",
                kol_name="测试",
            )
            from docx import Document
            doc = Document(path)
            texts = [p.text for p in doc.paragraphs]
            assert "Main" in texts
            assert "Sub" in texts
        finally:
            mod.REPORT_DIR = original_dir
