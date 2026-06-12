"""
Unit tests for persona_docx service.

覆盖：
- 正常路径：profile/plan 生成、问卷模板
- 边界条件：空内容、超长内容、特殊字符、多级标题、嵌套列表、引用、粗体
- Markdown 解析：标题 1-4、无序列表、有序列表、引用、粗体行内格式
"""
import os

import pytest

from app.services.persona_docx import (
    generate_persona_docx,
    generate_questionnaire_template,
    _markdown_to_paragraphs,
)
from docx import Document


# ── 正常路径 ─────────────────────────────────────────────────────


def test_generate_profile_docx():
    content = "# 测试达人 · 人格档案 v1.0\n\n## 一句话定位\n测试定位\n\n- **特点1**: 值1\n- **特点2**: 值2\n\n1. 第一条\n2. 第二条\n\n> 引用文本\n\n普通段落 **粗体** 内容。"
    path = generate_persona_docx(99991, "profile", content, "测试达人")
    assert os.path.exists(path)
    assert "99991_profile.docx" in path
    os.remove(path)


def test_generate_plan_docx():
    content = "# 测试达人 · 内容规划\n\n## 内容体系\n\n内容线1\n内容线2"
    path = generate_persona_docx(99992, "plan", content, "测试达人")
    assert os.path.exists(path)
    assert "99992_plan.docx" in path
    os.remove(path)


def test_generate_questionnaire_template():
    path = generate_questionnaire_template()
    assert os.path.exists(path)
    assert path.endswith("questionnaire_template.docx")
    os.remove(path)


# ── 边界条件 ─────────────────────────────────────────────────────


def test_generate_docx_with_empty_content():
    """空内容不应崩溃，应生成有效 docx。"""
    path = generate_persona_docx(99993, "profile", "", "空达人")
    assert os.path.exists(path)
    assert os.path.getsize(path) > 0
    os.remove(path)


def test_generate_docx_with_very_long_content():
    """超长内容（10000 字）不应崩溃。"""
    long_text = "\n\n".join([f"## 段落 {i}\n这是第{i}段内容，" + "很长" * 100 for i in range(200)])
    path = generate_persona_docx(99994, "profile", long_text, "超长达人")
    assert os.path.exists(path)
    assert os.path.getsize(path) > 1000
    os.remove(path)


def test_generate_docx_with_special_characters():
    """特殊字符（emoji、中文标点、HTML 实体）不应崩溃。"""
    content = "# 达人·测试 🔥\n\n## 特殊字符\n\n- 中文标点：，。！？；：\n- HTML 实体：&amp; &lt; &gt;\n- Emoji：🎉🚀💯\n\n> 引用含特殊字符「」【】"
    path = generate_persona_docx(99995, "profile", content, "特殊字符达人")
    assert os.path.exists(path)
    os.remove(path)


def test_generate_docx_with_no_heading():
    """纯文本无标题，不应崩溃。"""
    content = "这是一段纯文本，没有任何 Markdown 格式。\n\n第二段也是纯文本。"
    path = generate_persona_docx(99996, "profile", content, "纯文本达人")
    assert os.path.exists(path)
    doc = Document(path)
    # 应有文档标题 + 内容段落
    assert len(doc.paragraphs) >= 2
    os.remove(path)


# ── Markdown 解析验证 ────────────────────────────────────────────


def test_markdown_heading_levels():
    """验证 # ~ #### 都被正确解析为 heading 1-4。"""
    doc = Document()
    content = "# 一级标题\n## 二级标题\n### 三级标题\n#### 四级标题"
    _markdown_to_paragraphs(doc, content)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert len(headings) == 4
    assert headings[0].style.name == "Heading 1"
    assert headings[1].style.name == "Heading 2"
    assert headings[2].style.name == "Heading 3"
    assert headings[3].style.name == "Heading 4"


def test_markdown_unordered_list():
    """验证 - 和 * 开头的无序列表。"""
    doc = Document()
    content = "- 列表项1\n- 列表项2\n* 星号列表项"
    _markdown_to_paragraphs(doc, content)

    bullets = [p for p in doc.paragraphs if "List Bullet" in p.style.name]
    assert len(bullets) == 3


def test_markdown_ordered_list():
    """验证 1. 2. 3. 有序列表。"""
    doc = Document()
    content = "1. 第一项\n2. 第二项\n3. 第三项"
    _markdown_to_paragraphs(doc, content)

    numbered = [p for p in doc.paragraphs if "List Number" in p.style.name]
    assert len(numbered) == 3


def test_markdown_blockquote():
    """验证 > 引用格式（斜体 + 缩进）。"""
    doc = Document()
    content = "> 这是一段引用文本"
    _markdown_to_paragraphs(doc, content)

    quote_para = doc.paragraphs[-1]
    assert quote_para.runs[0].italic is True


def test_markdown_bold_inline():
    """验证 **粗体** 行内格式。"""
    doc = Document()
    content = "普通文本 **粗体部分** 普通文本"
    _markdown_to_paragraphs(doc, content)

    para = doc.paragraphs[-1]
    bold_runs = [r for r in para.runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "粗体部分"


def test_markdown_multiple_bold_in_line():
    """验证一行中有多个 **粗体** 片段。"""
    doc = Document()
    content = "**第一个** 中间 **第二个**"
    _markdown_to_paragraphs(doc, content)

    para = doc.paragraphs[-1]
    bold_runs = [r for r in para.runs if r.bold]
    assert len(bold_runs) == 2
    assert bold_runs[0].text == "第一个"
    assert bold_runs[1].text == "第二个"


def test_markdown_mixed_content():
    """验证混合格式：标题 + 列表 + 引用 + 粗体 + 空行。"""
    doc = Document()
    content = (
        "# 主标题\n\n"
        "## 子标题\n\n"
        "- **要点1**: 说明\n"
        "- **要点2**: 说明\n\n"
        "> 重要提示\n\n"
        "普通段落 **强调** 文本。\n\n"
        "1. 有序1\n"
        "2. 有序2"
    )
    _markdown_to_paragraphs(doc, content)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    bullets = [p for p in doc.paragraphs if "List Bullet" in p.style.name]
    numbered = [p for p in doc.paragraphs if "List Number" in p.style.name]

    assert len(headings) == 2
    assert len(bullets) == 2
    assert len(numbered) == 2


def test_markdown_empty_lines_skipped():
    """验证空行被正确跳过。"""
    doc = Document()
    content = "第一行\n\n\n\n第二行"
    _markdown_to_paragraphs(doc, content)

    # 空行不生成段落，只有两行内容
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    assert len(non_empty) == 2


def test_docx_file_is_valid():
    """生成的 docx 可以被 python-docx 重新打开。"""
    content = "# 标题\n\n- 列表项\n\n> 引用"
    path = generate_persona_docx(99997, "profile", content, "验证达人")
    assert os.path.exists(path)

    # 重新打开验证
    doc = Document(path)
    assert len(doc.paragraphs) > 0
    os.remove(path)
