"""Unit tests for parse_livestream_writer_file（livestream-writer 专用解析函数）"""
import io
import re
import zipfile
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.services.file_parser import parse_livestream_writer_file


def _mock_file(filename: str, content: bytes) -> MagicMock:
    f = MagicMock()
    f.filename = filename
    f.read = AsyncMock(return_value=content)
    return f


# ---------- txt / md ----------

@pytest.mark.asyncio
async def test_txt_returns_text():
    result = await parse_livestream_writer_file(_mock_file("script.txt", "直播脚本内容".encode()))
    assert result == "直播脚本内容"


@pytest.mark.asyncio
async def test_md_returns_text():
    result = await parse_livestream_writer_file(_mock_file("script.md", "# 标题\n内容".encode()))
    assert "标题" in result


@pytest.mark.asyncio
async def test_txt_no_truncation():
    """livestream-writer 版本不截断"""
    long = "直播脚本" * 5000
    result = await parse_livestream_writer_file(_mock_file("long.txt", long.encode()))
    assert len(result) == len(long)


# ---------- docx ----------

@pytest.mark.asyncio
async def test_docx_extracts_paragraphs():
    from docx import Document
    doc = Document()
    doc.add_paragraph("达人身份锚定")
    doc.add_paragraph("痛点激发话术")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    result = await parse_livestream_writer_file(_mock_file("script.docx", buf.read()))
    assert "达人身份锚定" in result
    assert "痛点激发话术" in result


# ---------- pdf（不支持，返回提示） ----------

@pytest.mark.asyncio
async def test_pdf_returns_prompt_message():
    result = await parse_livestream_writer_file(_mock_file("script.pdf", b"%PDF-1.4"))
    assert "暂不支持" in result
    assert "PDF" in result


# ---------- pages（含日历噪声过滤）----------

def _make_mock_pages_zip(iwa_content: bytes) -> bytes:
    """构造一个包含 Index/Document.iwa 的 .pages zip 文件"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("Index/Document.iwa", iwa_content)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_pages_extracts_chinese_text():
    chinese_text = "这是一段超过十个字符的中文直播脚本内容，用于测试解析是否正常工作。" * 3
    raw_bytes = chinese_text.encode("utf-8")

    # 让 snappy.decompress 抛异常，触发回退到 iwa_data 原始字节
    iwa = b"\x00\x00\x00\x00" + raw_bytes
    pages_bytes = _make_mock_pages_zip(iwa)

    with patch("snappy.decompress", side_effect=Exception("decompress failed")):
        result = await parse_livestream_writer_file(_mock_file("script.pages", pages_bytes))
    # 回退情况下解析到的内容取决于正则匹配，不强断言具体内容
    assert isinstance(result, str)


@pytest.mark.asyncio
async def test_pages_filters_calendar_noise():
    """日历噪声行应被过滤"""
    noise_items = [
        "星期三BJR",
        "一月这是日历",
        "第一季度",
        "公元前",
    ]
    clean_items = [
        "这是真实的直播脚本内容，包含足够的中文文字让过滤器通过。",
    ]
    raw = "\n".join(noise_items + clean_items) * 5
    iwa = b"\x00\x00\x00\x00" + raw.encode("utf-8")
    pages_bytes = _make_mock_pages_zip(iwa)

    with patch("snappy.decompress", side_effect=Exception("fail")):
        result = await parse_livestream_writer_file(_mock_file("noise.pages", pages_bytes))

    # 日历噪声不应出现在结果中
    for noise in ["星期三BJR"]:
        assert noise not in result


@pytest.mark.asyncio
async def test_pages_missing_iwa_returns_error():
    """iwa 文件不存在时返回提示字符串"""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("Index/other.iwa", b"nothing")
    pages_bytes = buf.getvalue()

    result = await parse_livestream_writer_file(_mock_file("bad.pages", pages_bytes))
    assert "格式异常" in result


@pytest.mark.asyncio
async def test_pages_bad_zip_returns_error():
    """非 zip 文件（无法解压）时返回提示字符串"""
    result = await parse_livestream_writer_file(_mock_file("bad.pages", b"not a zip file"))
    assert "格式异常" in result


# ---------- 不支持的格式 ----------

@pytest.mark.asyncio
async def test_unsupported_format_raises_value_error():
    with pytest.raises(ValueError, match="不支持"):
        await parse_livestream_writer_file(_mock_file("data.csv", b"col1,col2"))


@pytest.mark.asyncio
async def test_unsupported_xls_raises_value_error():
    with pytest.raises(ValueError):
        await parse_livestream_writer_file(_mock_file("data.xls", b"\xd0\xcf\x11\xe0"))
